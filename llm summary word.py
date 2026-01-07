import fitz  # PyMuPDF
import re
import spacy
from openpyxl import Workbook
from docx import Document
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM

# -------------------------
# Load NLP and lightweight LLM
# -------------------------
nlp = spacy.load("en_core_web_sm")
tokenizer = AutoTokenizer.from_pretrained("sshleifer/distilbart-cnn-12-6")
model = AutoModelForSeq2SeqLM.from_pretrained("sshleifer/distilbart-cnn-12-6")

# -------------------------
# Keywords & Noise Patterns
# -------------------------
keywords = [
    "subscriber", "BHCA", "Erlang", "busy hour", "data usage",
    "call attempts", "number of nodes", "core network", "throughput",
    "redundancy", "capacity", "scalability", "latency", "network", "interfaces",
    "connectivity", "roaming", "ipv6", "ipv4", "link", "protocol", "utilization",
    "statistics","impact","performance","capacity","scaling","measurement",
    "mtbf","bottlenecks","limit","bottleneck","mttr","limits","protect",
    "interception","segregation","lawful","access","authentication","zones",
    "security","authorization","integrity","bss","ems","oss","configuration",
    "management","inventory","vnfm","nfmf","administration","monitoring",
    "release","software","upgrade","update","version","functionality",
    "memory","physical","cpu","component","resources","orchestrator",
    "container","cnf","platform","nfv","vnf","nfvi","virtualization",
    "pnf","Kubernetes","validation","automation","policy","slicing",
    "mirroring","tracing","test","rules","openstack","etsi","cncf",
    "compliance","deviations","3gpp","5GNSA","4GNSA","5G SA","Cloud",
    "MicroServices","Canary","Orchestration","FCAPS","GUI","SBI","NAS",
    "IMS","TS","802.1p","DSCP","REST","N26","S5-C","S5-U","API","HTTP",
    "FQDN","oAUTH2.0","Interworking","VoNR","VoWiFi","ViNR","SMS","QoS",
    "DPI","ATSSS","EPC","NF","AMF","MME","SMF","SGW","UPF","NFs","NRF",
    "NSSF","SCP"
]

ignore_patterns = [
    r"copyright", r"confidential", r"all rights reserved",
    r"proprietary", r"©", r"do not distribute", r"page \d+ of \d+",
    r"Commercial in Confidence", r"^\d{4}-\d{2}-\d{2}",
    r"Rev\s+[A-Z0-9]+", r"^\d+\s*\(\d+\)$", r"^This solution.*"
]

heading_pattern = re.compile(r"^\d+(\.\d+)*\s+([A-Z][\w\s\-\&]+)", re.MULTILINE)

# -------------------------
# Helper Functions
# -------------------------
def is_noise(text):
    for patt in ignore_patterns:
        if re.search(patt, text, re.IGNORECASE):
            return True
    return False

def clean_page_text(text, top_lines=2, bottom_lines=2):
    lines = text.splitlines()
    if len(lines) > (top_lines + bottom_lines):
        lines = lines[top_lines:len(lines)-bottom_lines]
    text = "\n".join(lines)
    text = re.sub(r"\.{3,}", " ", text)  # remove TOC dots
    text = re.sub(r"\n\s*\n", "\n", text)
    text = re.sub(r"\s{2,}", " ", text)
    return text.strip()

def extract_text_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    pages = []
    for page_num, page in enumerate(doc, start=1):
        pages.append(page.get_text("text"))
    return pages

def extract_relevant_sentences(text, keyword_list):
    relevant_sentences = []
    doc = nlp(text)
    for sent in doc.sents:
        if is_noise(sent.text):
            continue
        for keyword in keyword_list:
            if re.search(rf"(?<!\w){re.escape(keyword)}(?!\w)", sent.text, re.IGNORECASE):
                relevant_sentences.append(sent.text.strip())
                break
    return relevant_sentences

# -------------------------
# Section-wise grouping based on headings
# -------------------------
def group_sentences_by_headings(text, keyword_list):
    sections = {}
    headings = list(heading_pattern.finditer(text))
    if not headings:
        # fallback to whole text as single section
        sections["Full Document"] = extract_relevant_sentences(text, keyword_list)
        return sections

    for i, heading in enumerate(headings):
        section_title = heading.group(0).strip()
        start = heading.end()
        end = headings[i+1].start() if i+1 < len(headings) else len(text)
        section_text = text[start:end].strip()
        sentences = extract_relevant_sentences(section_text, keyword_list)
        if sentences:
            sections[section_title] = sentences
    return sections

# -------------------------
# Chunked LLM Summarization
# -------------------------
def chunk_text(sentences, max_tokens=400):
    chunks = []
    current_chunk = ""
    token_count = 0
    for s in sentences:
        tokens = tokenizer.tokenize(s)
        if token_count + len(tokens) > max_tokens:
            chunks.append(current_chunk.strip())
            current_chunk = s
            token_count = len(tokens)
        else:
            current_chunk += " " + s
            token_count += len(tokens)
    if current_chunk:
        chunks.append(current_chunk.strip())
    return chunks

def summarize_with_llm(sentences, max_length=150):
    chunks = chunk_text(sentences, max_tokens=400)
    summaries = []
    for chunk in chunks:
        inputs = tokenizer(chunk, return_tensors="pt", truncation=True, max_length=1024)
        summary_ids = model.generate(inputs["input_ids"], max_length=max_length, min_length=30, length_penalty=2.0)
        summary_text = tokenizer.decode(summary_ids[0], skip_special_tokens=True)
        summaries.append(summary_text)
    return " ".join(summaries)

# -------------------------
# Export
# -------------------------
def save_summary_to_word(section_summaries, filename="section_wise_summary.docx"):
    doc = Document()
    doc.add_heading("Section-wise Summary", level=0)
    for sec, sentences in section_summaries.items():
        doc.add_heading(sec, level=1)
        summary = summarize_with_llm(sentences)
        doc.add_paragraph(summary)
    doc.save(filename)
    print(f"✅ Section-wise summary saved to {filename}")

def export_to_excel(data, filename="keyword_matches.xlsx"):
    wb = Workbook()
    ws = wb.active
    ws.title = "Matches"
    ws.append(["Page", "Keyword", "Sentence"])
    for item in data:
        ws.append([item["page"], item["keyword"], item["sentence"]])
    wb.save(filename)
    print(f"✅ Keyword matches exported to {filename}")

# -------------------------
# Main Process
# -------------------------
def process_rfp(pdf_path):
    pages = extract_text_from_pdf(pdf_path)
    full_text = "\n".join([clean_page_text(p) for p in pages])
    # Section-wise
    sections = group_sentences_by_headings(full_text, keywords)

    # Flatten all sentences for Excel
    relevant = []
    for sec, sentences in sections.items():
        for s in sentences:
            relevant.append({"page": "-", "keyword": "-", "sentence": s})

    export_to_excel(relevant, "keyword_matches.xlsx")
    save_summary_to_word(sections, "section_wise_summary.docx")
    print("\n✅ Process Completed!")

# -------------------------
# Run
# -------------------------
pdf_path = "C:/Users/eaprmis/Downloads/test.pdf"  # Replace with your file path
process_rfp(pdf_path)
